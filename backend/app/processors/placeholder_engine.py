"""
占位符解析引擎
支持 {{field}} 格式的占位符
支持段落和表格中的占位符
"""
import re
from typing import Dict, List, Any, Tuple, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class PlaceholderEngine:
    """占位符解析引擎"""

    # 匹配 {{field}} 格式的占位符（支持中文）
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([\w一-鿿]+)\}\}')

    def extract_placeholders(self, template_path: str) -> List[str]:
        """
        从Word模板中提取所有占位符
        返回占位符名称列表（去重）
        """
        doc = Document(template_path)
        placeholders = set()

        # 从段落中提取
        for paragraph in doc.paragraphs:
            placeholders.update(self._extract_from_text(paragraph.text))

        # 从表格中提取
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(self._extract_from_text(paragraph.text))

        return list(placeholders)

    def _extract_from_text(self, text: str) -> List[str]:
        """从文本中提取占位符"""
        if not text or '{{' not in text:
            return []
        matches = self.PLACEHOLDER_PATTERN.findall(text)
        return matches

    def fill_template(self, template_path: str, data: Dict[str, Any], output_path: str, formats: Dict[str, Dict] = None) -> bool:
        """
        填充模板
        template_path: 模板文件路径
        data: 填充数据 {"字段名": "值"}
        output_path: 输出文件路径
        formats: 格式设置 {"字段名": {"font_name": "宋体", "font_size": "五号", "bold": 1, "color": "#FF0000", "alignment": "center"}}
        """
        doc = Document(template_path)

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, formats)

        # 替换表格中的占位符
        for table in doc.tables:
            self._replace_in_table(table, data, formats)

        # 保存文档
        doc.save(output_path)

        # 清除文档默认样式中的主题字体引用，避免覆盖run级别的字体设置
        self._clear_document_theme_fonts(output_path)

        return True

    def _clear_document_theme_fonts(self, docx_path: str) -> None:
        """清除文档中的主题字体设置和段落级别的字体设置，直接修改ZIP内的XML"""
        import zipfile
        from lxml import etree

        # 读取ZIP
        with zipfile.ZipFile(docx_path, 'r') as zin:
            content = zin.namelist()
            data = {}
            for name in content:
                data[name] = zin.read(name)

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # 修改 styles.xml - 清除 docDefaults 中的主题字体
        if 'word/styles.xml' in data:
            styles_xml = data['word/styles.xml']
            root = etree.fromstring(styles_xml)

            # 找到 docDefaults
            doc_defaults = root.find('.//w:docDefaults', ns)
            if doc_defaults is not None:
                rpr_default = doc_defaults.find('.//w:rPrDefault/w:rPr', ns)
                if rpr_default is not None:
                    rfonts = rpr_default.find('w:rFonts', ns)
                    if rfonts is not None:
                        # 删除所有主题相关属性
                        theme_attrs = []
                        for attr in list(rfonts.attrib.keys()):
                            attr_lower = attr.lower()
                            # 检查是否是主题属性（包括带命名空间的）
                            if 'theme' in attr_lower or 'cstheme' in attr_lower:
                                theme_attrs.append(attr)
                        for attr in theme_attrs:
                            del rfonts.attrib[attr]
                        # 如果删除后 rFonts 变成空元素，直接删除它
                        if len(rfonts.attrib) == 0:
                            rpr_default.remove(rfonts)

            data['word/styles.xml'] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')

        # 修改 document.xml - 清除段落级别的 rPr 字体设置
        if 'word/document.xml' in data:
            doc_xml = data['word/document.xml']
            root = etree.fromstring(doc_xml)

            # 清除段落/表格级别的 rFonts（不是 run 级别）
            for pPr in root.findall('.//w:pPr', ns):
                rPr = pPr.find('w:rPr', ns)
                if rPr is not None:
                    rfonts = rPr.find('w:rFonts', ns)
                    if rfonts is not None:
                        rPr.remove(rfonts)

            data['word/document.xml'] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')

        # 写回文件
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name in content:
                zout.writestr(name, data[name])

    def _replace_in_paragraph(self, paragraph: Paragraph, data: Dict[str, Any], formats: Dict[str, Dict] = None) -> None:
        """替换段落中的占位符，应用格式设置"""
        if not paragraph.text or '{{' not in paragraph.text:
            return

        # 收集需要处理的run
        runs_to_process = []
        for run in paragraph.runs:
            if '{{' not in run.text:
                continue
            runs_to_process.append(run)

        for run in runs_to_process:
            original_text = run.text
            matches = list(self.PLACEHOLDER_PATTERN.finditer(original_text))

            if not matches:
                continue

            # 检查是否有需要替换的占位符
            has_replacement = any(m.group(1) in data for m in matches)
            if not has_replacement:
                continue

            # 只有一个占位符且没有普通文本，直接替换
            if len(matches) == 1 and original_text == matches[0].group(0):
                field_name = matches[0].group(1)
                if field_name in data:
                    value = str(data[field_name]) if data[field_name] is not None else ''
                    run.text = value
                    self._apply_format_to_run(run, field_name, formats)
                    if formats and field_name in formats:
                        fmt = formats[field_name]
                        align_map = {'left': 0, 'center': 1, 'right': 2, 'justify': 3}
                        if fmt.get('alignment') in align_map:
                            paragraph.alignment = align_map[fmt['alignment']]
                continue

            # 多个占位符或有普通文本，需要拆分重建
            # 构建片段列表：(文本, 字段名或None)
            segments = []
            last_end = 0
            for match in matches:
                if match.start() > last_end:
                    segments.append((original_text[last_end:match.start()], None))
                segments.append((match.group(0), match.group(1)))
                last_end = match.end()
            if last_end < len(original_text):
                segments.append((original_text[last_end:], None))

            # 获取原始run的XML元素
            run_elem = run._element
            p_elem = paragraph._p

            # 保存原始run的位置索引
            run_idx = list(p_elem).index(run_elem)

            # 保存原始rPr用于非占位符片段
            original_rpr = run_elem.find(qn('w:rPr'))
            original_style = None
            if original_rpr is not None:
                style_elem = original_rpr.find(qn('w:rStyle'))
                if style_elem is not None:
                    original_style = style_elem.get(qn('w:val'))

            # 克隆原始run元素用于创建新run
            def clone_run_with_text(text, field_name=None):
                """创建新run元素"""
                from docx.oxml import OxmlElement
                from docx.text.run import Run

                new_r = OxmlElement('w:r')
                # 复制原始rPr（如果有）
                if original_rpr is not None:
                    new_rpr = OxmlElement('w:rPr')
                    for child in original_rpr:
                        new_rpr.append(child)
                    new_r.append(new_rpr)

                # 添加文本
                new_t = OxmlElement('w:t')
                new_t.text = text
                new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_r.append(new_t)

                return new_r

            # 重建段落：删除原run，用新片段替换
            # 先保存原始run的位置信息
            run_parent = run_elem.getparent()
            run_position = list(run_parent).index(run_elem)

            # 移除原始run
            run_elem.addnext(run_elem)  # This doesn't do what I want, let me think...

            # Actually, let me use a different approach: modify in place
            # For each segment, create new runs and insert them properly

            # First, let's take a simpler approach:
            # 1. Save all segment info
            # 2. Clear the original run
            # 3. Insert new runs in correct order

            # Clear the original run's text
            run.text = ''

            # Build list of new run elements
            new_runs = []
            for text, field_name in segments:
                if not text:
                    continue

                if field_name is not None and field_name in data:
                    # Placeholder - create run with value
                    value = str(data[field_name]) if data[field_name] is not None else ''
                    new_r = clone_run_with_text(value, field_name)

                    # Apply formatting if there's a field_name with format
                    if formats and field_name in formats:
                        fmt = formats[field_name]
                        rpr_new = new_r.find(qn('w:rPr'))
                        if rpr_new is None:
                            rpr_new = OxmlElement('w:rPr')
                            new_r.insert(0, rpr_new)

                        # Remove style if present
                        style_elem = rpr_new.find(qn('w:rStyle'))
                        if style_elem is not None:
                            rpr_new.remove(style_elem)

                        # Add fonts
                        rfonts = rpr_new.find(qn('w:rFonts'))
                        if rfonts is None:
                            rfonts = OxmlElement('w:rFonts')
                            rpr_new.insert(0, rfonts)

                        if fmt.get('font_name'):
                            font_name = fmt['font_name']
                            rfonts.set(qn('w:eastAsia'), font_name)
                            rfonts.set(qn('w:ascii'), font_name)
                            rfonts.set(qn('w:hAnsi'), font_name)
                            rfonts.set(qn('w:cs'), font_name)
                            rfonts.set(qn('w:eastAsiaAscii'), font_name)

                        # Add font size
                        if fmt.get('font_size'):
                            size_map = {'八号': 5, '七号': 8, '六号': 8.7, '五号': 10.5, '小五': 9, '一号': 26, '小一': 24, '二号': 22, '小二': 18, '三号': 16, '小三': 15, '四号': 14, '小四': 12, '五号': 10.5}
                            size = size_map.get(fmt['font_size'], 10.5)
                            sz = OxmlElement('w:sz')
                            sz.set(qn('w:val'), str(int(size * 2)))
                            rpr_new.append(sz)
                            szCs = OxmlElement('w:szCs')
                            szCs.set(qn('w:val'), str(int(size * 2)))
                            rpr_new.append(szCs)

                        # Add bold
                        if fmt.get('bold'):
                            b = OxmlElement('w:b')
                            rpr_new.append(b)

                        # Add italic
                        if fmt.get('italic'):
                            i = OxmlElement('w:i')
                            rpr_new.append(i)

                        # Add color
                        if fmt.get('color'):
                            color = fmt['color'].replace('#', '')
                            if len(color) == 6:
                                color_elem = OxmlElement('w:color')
                                color_elem.set(qn('w:val'), color)
                                rpr_new.append(color_elem)

                    new_runs.append((new_r, field_name))
                else:
                    # Non-placeholder text - clone original style
                    new_r = clone_run_with_text(text, None)
                    new_runs.append((new_r, None))

            # Now insert new runs in correct order (right to left)
            # This way, each insert doesn't affect the position of previously inserted runs
            insert_pos = run_elem
            for new_r, field_name in reversed(new_runs):
                insert_pos.addnext(new_r)

            # Now remove the empty original run
            run_elem.getparent().remove(run_elem)

            # Handle alignment - use first placeholder's alignment
            if formats:
                for _, field_name in new_runs:
                    if field_name and field_name in formats:
                        fmt = formats[field_name]
                        align_map = {'left': 0, 'center': 1, 'right': 2, 'justify': 3}
                        if fmt.get('alignment') in align_map:
                            paragraph.alignment = align_map[fmt['alignment']]
                        break

    def _apply_format_to_run(self, run: Paragraph, field_name: str, formats: Dict[str, Dict] = None) -> None:
        """应用格式设置到run"""
        if not formats or field_name not in formats:
            return

        fmt = formats[field_name]

        # 获取或创建 rPr 元素
        rpr = run._element.get_or_add_rPr()

        # 清除可能存在的 rStyle 引用，确保不受样式影响
        style_elem = rpr.find(qn('w:rStyle'))
        if style_elem is not None:
            rpr.remove(style_elem)

        # 清除可能存在的样式继承 - 设置 rFonts
        rfonts = rpr.get_or_add_rFonts()
        if fmt.get('font_name'):
            font_name = fmt['font_name']
            # 设置所有字体属性，覆盖默认样式
            rfonts.set(qn('w:eastAsia'), font_name)
            rfonts.set(qn('w:ascii'), font_name)
            rfonts.set(qn('w:hAnsi'), font_name)
            rfonts.set(qn('w:cs'), font_name)
            rfonts.set(qn('w:eastAsiaAscii'), font_name)

        if fmt.get('font_size'):
            size_map = {'八号': 5, '七号': 8, '六号': 8.7, '五号': 10.5, '小五': 9, '一号': 26, '小一': 24, '二号': 22, '小二': 18, '三号': 16, '小三': 15, '四号': 14, '小四': 12, '五号': 10.5}
            size = size_map.get(fmt['font_size'], 10.5)
            run.font.size = Pt(size)

        if fmt.get('bold'):
            run.font.bold = True
        if fmt.get('italic'):
            run.font.italic = True
        if fmt.get('color'):
            color = fmt['color'].replace('#', '')
            if len(color) == 6:
                r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)

    def _replace_in_table(self, table: Table, data: Dict[str, Any], formats: Dict[str, Dict] = None) -> None:
        """替换表格中的占位符"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, data, formats)

    def mark_field(self, template_path: str, original_text: str, field_name: str, output_path: str = None) -> bool:
        """
        将文档中的指定文本替换为占位符 {{field_name}}
        template_path: 原始模板路径
        original_text: 要替换的原文本
        field_name: 字段名
        output_path: 输出路径（如果为None，则覆盖原文件）
        返回: 是否成功
        """
        if not original_text or not field_name:
            return False

        placeholder = "{{" + field_name + "}}"
        doc = Document(template_path)
        replaced = False

        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            if self._replace_text_in_paragraph(paragraph, original_text, placeholder):
                replaced = True

        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._replace_text_in_paragraph(paragraph, original_text, placeholder):
                            replaced = True

        # 保存文件（无论是否发生替换）
        save_path = output_path if output_path else template_path
        doc.save(save_path)

        return replaced

    def insert_field(self, template_path: str, field_name: str, output_path: str = None,
                  position: str = "end", table_index: Optional[int] = None, row: Optional[int] = None, col: Optional[int] = None) -> bool:
        """
        在文档中插入新的占位符字段
        template_path: 模板路径
        field_name: 字段名
        output_path: 输出路径
        position: "end" 在文档末尾添加 | "start" 在开头添加 | "table" 在表格单元格中添加
        table_index: 表格索引（第几个表格，0开始）
        row: 行索引（0开始）
        col: 列索引（0开始）
        返回: 是否成功
        """
        if not field_name:
            return False

        placeholder = "{{" + field_name + "}}"
        doc = Document(template_path)

        if position == "table" and table_index is not None and row is not None and col is not None:
            # 在指定表格的指定单元格中插入
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                if row < len(table.rows) and col < len(table.rows[row].cells):
                    cell = table.rows[row].cells[col]
                    # 清除单元格内容并添加占位符
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
                    if cell.paragraphs:
                        cell.paragraphs[0].text = placeholder
                    else:
                        cell.add_paragraph(placeholder)
                    save_path = output_path if output_path else template_path
                    doc.save(save_path)
                    return True
            return False
        elif position == "start" and doc.paragraphs:
            # 在第一个段落开头插入 - 使用XML方式插入
            from docx.oxml import OxmlElement
            first_para = doc.paragraphs[0]
            # 创建一个新的run元素
            new_run = OxmlElement('w:r')
            # 添加文本
            new_t = OxmlElement('w:t')
            new_t.text = placeholder
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_run.append(new_t)
            # 在段落的第一个位置插入
            first_para._p.insert(0, new_run)
            save_path = output_path if output_path else template_path
            doc.save(save_path)
            return True
        else:
            # 在文档末尾添加新段落
            doc.add_paragraph(placeholder)
            save_path = output_path if output_path else template_path
            doc.save(save_path)
            return True

    def _replace_text_in_paragraph(self, paragraph: Paragraph, original_text: str, placeholder: str) -> bool:
        """
        将段落中的指定文本替换为占位符
        规则：
        1. 如果 original_text 同时存在于标签和 {{}} 内的占位符中，跳过（已有占位符）
        2. 如果 original_text 只在标签中，替换标签
        3. 如果 original_text 只在占位符中，无操作
        """
        if not original_text or original_text not in paragraph.text:
            return False

        import re
        # 检查是否有受保护的占位符（包含在 {{}} 内）
        protected_pattern = re.compile(r'\{\{[^}]*' + re.escape(original_text) + r'[^}]*\}\}')

        full_text = paragraph.text
        protected_matches = list(protected_pattern.finditer(full_text))

        # 如果 original_text 已经在占位符中，且也在标签中，则跳过（标签已有关联的占位符）
        if protected_matches:
            # 检查标签中是否也有 original_text
            # 找到第一个受保护区域
            first_protected_start = protected_matches[0].start()
            # 检查 original_text 在受保护区域之前是否以标签形式出现
            label_pos = full_text.find(original_text)
            if label_pos >= 0 and label_pos < first_protected_start:
                # 标签在占位符之前，说明标签和占位符都包含 original_text
                # 这是用户试图标记已有关联的情况，跳过
                return False

        replaced = False
        runs = list(paragraph.runs)

        for run in runs:
            if original_text not in run.text:
                continue

            run_elem = run._element
            rpr = run_elem.find(qn('w:rPr'))
            has_rpr = rpr is not None

            # 如果run文本完全匹配，直接替换
            if run.text == original_text:
                run.text = placeholder
                replaced = True
            else:
                # 分割替换：处理多处出现的情况
                parts = run.text.split(original_text)
                if len(parts) >= 2:
                    # 用占位符连接所有部分
                    # 例如 "问题来源：xxx，问题来源：yyy".split("问题来源") = ["", "：xxx，", "：yyy"]
                    # 连接后变成 "{{问题来源}}：xxx，{{问题来源}}：yyy"
                    run.text = placeholder.join(parts)
                    replaced = True

        return replaced

    def validate_placeholders(self, template_path: str, required_fields: List[str]) -> Dict[str, Any]:
        """
        验证模板中的占位符是否满足要求
        """
        placeholders = set(self.extract_placeholders(template_path))
        required = set(required_fields)

        missing = required - placeholders
        extra = placeholders - required

        return {
            'valid': len(missing) == 0,
            'all_placeholders': list(placeholders),
            'missing_fields': list(missing),
            'extra_fields': list(extra)
        }

    def _merge_adjacent_runs(self, paragraph: Paragraph) -> None:
        """
        合并段落中相邻的 run（解决 Word 将中文/数字分成多个 run 的问题）
        如果两个相邻 run 的格式（字体、大小、颜色等）相同，则合并它们
        """
        runs = paragraph.runs
        if len(runs) <= 1:
            return

        # 从后向前合并，避免索引变化问题
        i = len(runs) - 1
        while i > 0:
            current_run = runs[i]
            prev_run = runs[i - 1]

            if self._runs_have_same_format(current_run, prev_run):
                # 合并文本：将前一个 run 的文本追加到当前 run
                prev_text = prev_run.text
                current_text = current_run.text

                # 先设置当前 run 的文本为空，再删除它
                current_run.text = ''
                prev_run.text = prev_text + current_text

                # 删除当前 run 的 XML 元素
                current_run._element.getparent().remove(current_run._element)
            else:
                i -= 1

    def _runs_have_same_format(self, run1, run2) -> bool:
        """
        比较两个 run 的格式是否相同
        """
        def get_run_style(run):
            """获取 run 的关键格式属性"""
            rpr = run._element.find(qn('w:rPr'))
            if rpr is None:
                return {}

            style = {}
            # 字体
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is not None:
                for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs', 'w:eastAsiaAscii']:
                    val = rfonts.get(qn(attr))
                    if val:
                        style[attr] = val

            # 字体大小
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                style['w:sz'] = sz.get(qn('w:val'))

            # 粗体
            b = rpr.find(qn('w:b'))
            if b is not None:
                style['bold'] = True

            # 斜体
            i = rpr.find(qn('w:i'))
            if i is not None:
                style['italic'] = True

            # 颜色
            color = rpr.find(qn('w:color'))
            if color is not None:
                style['color'] = color.get(qn('w:val'))

            return style

        style1 = get_run_style(run1)
        style2 = get_run_style(run2)

        return style1 == style2
