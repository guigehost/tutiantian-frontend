import pytest
import tempfile
import os
from pathlib import Path
from docx import Document
import pandas as pd

# 添加项目路径
import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from app.processors import (
    ExcelProcessor,
    WordProcessor,
    PlaceholderEngine,
    FileValidator
)

class TestPlaceholderEngine:
    """测试占位符引擎"""

    @pytest.fixture
    def template_path(self, tmp_path):
        """创建测试模板"""
        doc = Document()

        # 添加段落
        p = doc.add_paragraph('问题来源：{{问题来源}}')
        p = doc.add_paragraph('问题概述：{{问题概述}}')

        # 添加表格
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = '问题来源'
        table.cell(0, 1).text = '{{问题来源}}'
        table.cell(1, 0).text = '问题概述'
        table.cell(1, 1).text = '{{问题概述}}'
        table.cell(2, 0).text = '了解内容'
        table.cell(2, 1).text = '{{了解内容}}'

        path = tmp_path / "test_template.docx"
        doc.save(str(path))
        return str(path)

    def test_extract_placeholders(self, template_path):
        engine = PlaceholderEngine()
        placeholders = engine.extract_placeholders(template_path)
        assert '问题来源' in placeholders
        assert '问题概述' in placeholders
        assert '了解内容' in placeholders

    def test_fill_template(self, template_path, tmp_path):
        engine = PlaceholderEngine()
        output_path = tmp_path / "output.docx"

        data = {
            '问题来源': '单位A',
            '问题概述': '财务不规范',
            '了解内容': '需要进一步调查'
        }

        engine.fill_template(template_path, data, str(output_path))

        assert output_path.exists()

        # 验证填充结果
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert '单位A' in text
        assert '财务不规范' in text

    def test_validate_placeholders(self, template_path):
        engine = PlaceholderEngine()
        required_fields = ['问题来源', '问题概述', '了解内容']
        result = engine.validate_placeholders(template_path, required_fields)
        assert result['valid'] is True
        assert len(result['missing_fields']) == 0


class TestExcelProcessor:
    """测试Excel处理器"""

    @pytest.fixture
    def excel_path(self, tmp_path):
        """创建测试Excel"""
        df = pd.DataFrame({
            '问题来源': ['单位A', '单位B'],
            '问题概述': ['问题1', '问题2'],
            '了解内容': ['内容1', '内容2']
        })
        path = tmp_path / "test_data.xlsx"
        df.to_excel(str(path), index=False)
        return str(path)

    def test_read_headers(self, excel_path):
        processor = ExcelProcessor(excel_path)
        headers = processor.read_headers()
        assert '问题来源' in headers
        assert '问题概述' in headers

    def test_read_data(self, excel_path):
        processor = ExcelProcessor(excel_path)
        df = processor.read_data()
        assert len(df) == 2
        assert df.iloc[0]['问题来源'] == '单位A'

    def test_get_row_count(self, excel_path):
        processor = ExcelProcessor(excel_path)
        count = processor.get_row_count()
        assert count == 2

    def test_validate_columns(self, excel_path):
        processor = ExcelProcessor(excel_path)
        result = processor.validate_columns(['问题来源', '问题概述'])
        assert result['valid'] is True


class TestWordProcessor:
    """测试Word处理器"""

    def test_validate_format(self):
        processor = WordProcessor()
        assert processor.validate_format('test.docx') is True
        assert processor.validate_format('test.xlsx') is False


class TestFileValidator:
    """测试文件验证器"""

    def test_validate_excel_format(self, tmp_path):
        validator = FileValidator()
        # 创建测试Excel文件
        df = pd.DataFrame({'col1': [1, 2]})
        excel_path = tmp_path / "test.xlsx"
        df.to_excel(str(excel_path), index=False)

        result = validator.validate_excel(str(excel_path))
        assert result['valid'] is True

    def test_validate_word_format(self, tmp_path):
        validator = FileValidator()
        # 创建测试Word文件
        doc = Document()
        doc.add_paragraph('Test')
        word_path = tmp_path / "test.docx"
        doc.save(str(word_path))

        result = validator.validate_word(str(word_path))
        assert result['valid'] is True
        assert result.get('malware_check', {}).get('safe') is True

    def test_invalid_format(self):
        validator = FileValidator()
        result = validator.validate_excel('test.txt')
        assert result['valid'] is False
        assert '不支持的文件格式' in result['error']


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
